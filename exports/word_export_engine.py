"""
============================================================
XUẤT BẢN WORD - BỘ ĐIỀU PHỐI TRUNG TÂM (WORD EXPORT ENGINE)
Hỗ trợ Render Bảng (Tables), Markdown, và bảo toàn LaTeX.
============================================================
"""

import io
import re
import os
from loguru import logger
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordExportEngine:
    @staticmethod
    def _parse_inline_text(paragraph, text: str):
        """Hàm xử lý bôi đậm, in nghiêng và giữ nguyên LaTeX ($) trong từng đoạn hoặc ô bảng."""
        # Tách chuỗi dựa trên ký hiệu in đậm (**) và LaTeX ($$, $)
        tokens = re.split(r'(\*\*.*?\*\*|\$\$.*?\$\$|\$.*?\$)', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('$$') and token.endswith('$$'):
                run = paragraph.add_run(token)
                run.italic = True
            elif token.startswith('$') and token.endswith('$'):
                run = paragraph.add_run(token)
                run.italic = True
            else:
                paragraph.add_run(token)

    @staticmethod
    def convert_markdown_to_docx_bytes(markdown_text: str, template_path: str = None) -> bytes:
        try:
            if template_path and os.path.exists(template_path):
                doc = Document(template_path)
            else:
                doc = Document()
            
            # Căn lề chuẩn
            for section in doc.sections:
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(1.5)
            
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(13)
            
            if not markdown_text:
                markdown_text = "Không có nội dung xuất bản."

            lines = str(markdown_text).split('\n')
            
            in_table = False
            table_data = []

            def render_buffered_table():
                """Vẽ bảng từ mảng dữ liệu table_data đã thu thập"""
                if not table_data: return
                
                # Bỏ qua dòng phân cách của Markdown (VD: |---|---|)
                valid_rows = [row for row in table_data if not re.match(r'^[\s\|:.-]+$', row)]
                if not valid_rows: return
                
                # Tính số cột dựa trên dòng đầu tiên
                num_cols = len([c for c in valid_rows[0].strip().strip('|').split('|')])
                table = doc.add_table(rows=len(valid_rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for i, row_str in enumerate(valid_rows):
                    cells = [c.strip() for c in row_str.strip().strip('|').split('|')]
                    for j, cell_text in enumerate(cells):
                        if j < num_cols:
                            cell = table.cell(i, j)
                            cell.text = "" # Xóa nội dung mặc định
                            p = cell.paragraphs[0]
                            WordExportEngine._parse_inline_text(p, cell_text)
            
            for line in lines:
                line_clean = line.strip()
                
                # Phát hiện Bảng Markdown (bắt đầu và kết thúc bằng |)
                if line_clean.startswith('|') and line_clean.endswith('|'):
                    in_table = True
                    table_data.append(line_clean)
                    continue
                else:
                    # Nếu đang ở trong bảng mà gặp dòng chữ thường -> Kết thúc bảng, vẽ bảng ra
                    if in_table:
                        render_buffered_table()
                        in_table = False
                        table_data = []

                if not line_clean:
                    continue
                
                # Xử lý Tiêu đề (Headings)
                if line_clean.startswith('#'):
                    level = len(line_clean) - len(line_clean.lstrip('#'))
                    text = line_clean.lstrip('#').strip()
                    p = doc.add_heading(text, level=min(level, 3))
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        run.font.color.rgb = None 
                        run.font.size = Pt(16) if level == 1 else Pt(14)
                    continue

                # Xử lý đoạn văn bình thường (Paragraphs & Lists)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                if line_clean.startswith('*'):
                    p.paragraph_format.left_indent = Cm(0.5)
                    line_clean = line_clean.lstrip('* ')
                elif re.match(r'^[a-zA-Z]\)\s+', line_clean) or re.match(r'^\d+\.\s+', line_clean) or line_clean.startswith('-'):
                    p.paragraph_format.left_indent = Cm(0.3)
                else:
                    p.paragraph_format.first_line_indent = Cm(1.0)
                    
                WordExportEngine._parse_inline_text(p, line_clean)

            # Vẽ bảng cuối cùng nếu file kết thúc bằng một bảng
            if in_table:
                render_buffered_table()

            f = io.BytesIO()
            doc.save(f)
            return f.getvalue()
            
        except Exception as e:
            logger.error(f"Lỗi xuất bản file Word: {str(e)}")
            err_doc = Document()
            err_doc.add_paragraph(f"Đã xảy ra lỗi khi tạo file Word: {str(e)}")
            f = io.BytesIO()
            err_doc.save(f)
            return f.getvalue()
