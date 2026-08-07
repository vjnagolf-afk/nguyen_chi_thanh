import streamlit as st
import pandas as pd
from docx import Document
import io

# Cấu hình trang
st.set_page_config(page_title="Streamlit Test App", layout="centered")

# Tiêu đề ứng dụng
st.title("🚀 Chạy Thử Nghiệm Dự Án Streamlit")
st.write("Ứng dụng này kiểm tra xem các thư viện trong `requirements.txt` đã hoạt động tốt chưa.")

st.divider()

# --- Kiểm tra pandas ---
st.header("1. Kiểm tra Pandas")
st.write("Hiển thị một bảng dữ liệu (DataFrame) đơn giản:")

# Tạo dữ liệu mẫu
data = {
    "Tên sản phẩm": ["Sản phẩm A", "Sản phẩm B", "Sản phẩm C"],
    "Số lượng": [15, 30, 45],
    "Giá bán (VNĐ)": [100000, 250000, 150000]
}
df = pd.DataFrame(data)
st.dataframe(df, use_container_width=True)

st.divider()

# --- Kiểm tra python-docx ---
st.header("2. Kiểm tra Python-Docx")
st.write("Tạo và tải xuống một file Word cơ bản.")

def create_sample_word():
    # Khởi tạo document
    doc = Document()
    doc.add_heading('Tài Liệu Mẫu (Test Document)', 0)
    doc.add_paragraph('Đây là một tệp Word được tạo tự động thông qua thư viện python-docx trên Streamlit.')
    
    # Thêm một bảng vào Word dựa trên dữ liệu Pandas ở trên
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên sản phẩm'
    hdr_cells[1].text = 'Số lượng'
    hdr_cells[2].text = 'Giá bán (VNĐ)'
    
    for i in range(len(df)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(df.loc[i, "Tên sản phẩm"])
        row_cells[1].text = str(df.loc[i, "Số lượng"])
        row_cells[2].text = str(df.loc[i, "Giá bán (VNĐ)"])
        
    # Lưu file vào bộ đệm (BytesIO) để tải xuống
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Nút tạo và tải file Word
word_file = create_sample_word()
st.download_button(
    label="📄 Tải xuống file Word mẫu",
    data=word_file,
    file_name="tai_lieu_thu_nghiem.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
